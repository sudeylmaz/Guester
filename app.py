# -*- coding: utf-8 -*-
#
# Adı Soyadı: Ceydanur Arslan
# Öğrenci Numarası: 262484066
#
# Adı Soyadı: Sude Yılmaz
# Öğrenci Numarası: 262484068
#
from flask import Flask, render_template, request, redirect, session, jsonify
import os
import guester_ocr
import sqlite3
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.header import Header
import pdfplumber
import docx
import re
import socket
import subprocess
import threading
import atexit

# .env dosyasını oku ve çevre değişkenlerine ekle
env_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".env")
if os.path.exists(env_path):
    try:
        with open(env_path, "r", encoding="utf-8") as f:
            for line in f:
                if line.strip() and not line.startswith("#") and "=" in line:
                    key, val = line.strip().split("=", 1)
                    os.environ[key.strip()] = val.strip()
    except Exception as e:
        print(f"Guester: .env dosyası okunurken hata: {e}")

app = Flask(__name__)
app.secret_key = "guester123"

UPLOAD_FOLDER = "uploads"
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

g_public_url = None
g_tunnel_process = None
URL_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tunnel_url.txt")

def is_ssh_running():
    try:
        import platform
        if platform.system() == "Windows":
            out = subprocess.run(["tasklist"], capture_output=True, text=True, timeout=2)
            return "ssh.exe" in out.stdout.lower()
    except:
        pass
    return False

def cleanup_tunnel():
    global g_tunnel_process
    if g_tunnel_process:
        print("Guester: Tünel süreci sonlandırılıyor...")
        try:
            g_tunnel_process.terminate()
            g_tunnel_process.wait(timeout=2)
        except:
            pass

atexit.register(cleanup_tunnel)

def start_ssh_tunnel():
    global g_public_url, g_tunnel_process
    
    # Eğer SSH zaten çalışıyorsa ve dosyamızda link kayıtlıysa, yeniden tünel açma!
    if is_ssh_running() and os.path.exists(URL_FILE):
        try:
            with open(URL_FILE, "r") as f:
                saved_url = f.read().strip()
                if saved_url:
                    g_public_url = saved_url
                    print(f"Guester: Mevcut aktif tünel algılandı, kullanılıyor! -> {g_public_url}")
                    return
        except:
            pass

    print("Guester: İnternet Paylaşım Tüneli (localhost.run) başlatılıyor...")
    
    try:
        import platform
        if platform.system() == "Windows":
            subprocess.run(["taskkill", "/f", "/im", "ssh.exe"], capture_output=True, timeout=2)
    except:
        pass
        
    try:
        # Eski tünel dosyasını temizle
        if os.path.exists(URL_FILE):
            try: os.remove(URL_FILE)
            except: pass
            
        cmd = ["ssh", "-o", "StrictHostKeyChecking=no", "-R", "80:127.0.0.1:5000", "nokey@localhost.run"]
        process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1)
        g_tunnel_process = process
        for line in iter(process.stdout.readline, ''):
            print(f"TUNNEL: {line.strip()}")
            match = re.search(r'https://[a-zA-Z0-9.-]+\.lhr\.life', line)
            if match:
                g_public_url = match.group(0)
                print(f"Guester: İnternet Paylaşım Linkiniz hazır! -> {g_public_url}")
                # Dosyaya kaydet
                try:
                    with open(URL_FILE, "w") as f:
                        f.write(g_public_url)
                except:
                    pass
        process.stdout.close()
        process.wait()
    except Exception as e:
        print(f"Guester: Tünel başlatılamadı: {e}")

print("Guester baslatiliyor, akilli OCR sistemi devrede!...")

DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "guester.db")
OLD_DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "hosteasy.db")

if not os.path.exists(DB_PATH) and os.path.exists(OLD_DB_PATH):
    try:
        import shutil
        shutil.copy(OLD_DB_PATH, DB_PATH)
        print("Guester: Eski veritabanı (hosteasy.db) başarıyla guester.db olarak kopyalandı.")
    except Exception as e:
        print(f"Guester: Eski veritabanı kopyalanırken hata oluştu: {e}")

def db():
    return sqlite3.connect(DB_PATH)

def generate_unique_token():
    import string
    import random
    characters = string.ascii_letters
    while True:
        token = "".join(random.choices(characters, k=10))
        con = db()
        cur = con.cursor()
        cur.execute("SELECT id FROM davetiyeler WHERE davetiye_kod = ?", (token,))
        exists = cur.fetchone()
        con.close()
        if not exists:
            return token

def init_db():
    con = db()
    cur = con.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS kullanicilar(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ad_soyad TEXT,
            sifre TEXT,
            email TEXT UNIQUE
        )
    """)
    try:
        cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_kullanicilar_email ON kullanicilar(email)")
    except Exception as e:
        print(f"Guester: Benzersiz e-posta indeksi oluşturulurken hata: {e}")
    cur.execute("""
        CREATE TABLE IF NOT EXISTS davetiyeler(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            gelin_adi TEXT,
            damat_adi TEXT,
            saat_bilgisi TEXT,
            yer_bilgisi TEXT,
            giris_sifresi TEXT,
            davetiye_kod TEXT UNIQUE,
            olusturma_tarihi TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS konuklar(
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            davetiye_id INTEGER,
            isim TEXT,
            masa TEXT,
            durum TEXT DEFAULT 'Gelmedi',
            FOREIGN KEY(davetiye_id) REFERENCES davetiyeler(id)
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS sifre_kodlari(
            email TEXT PRIMARY KEY,
            kod TEXT,
            olusturma_tarihi TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)
    
    # Veritabanı göçü: davetiye_kod kolonu kontrolü
    cur.execute("PRAGMA table_info(davetiyeler)")
    columns = [c[1] for c in cur.fetchall()]
    
    if "davetiye_kod" not in columns:
        try:
            print("Hosteasy: davetiye_kod kolonu ekleniyor...")
            cur.execute("ALTER TABLE davetiyeler ADD COLUMN davetiye_kod TEXT")
            con.commit()
        except Exception as e:
            print(f"Hata (davetiye_kod ekleme): {e}")
            
    # Boş davetiye_kod'u olan kayıtları güncelle
    cur.execute("SELECT id FROM davetiyeler WHERE davetiye_kod IS NULL OR davetiye_kod = ''")
    rows_to_update = cur.fetchall()
    if rows_to_update:
        import string
        import random
        characters = string.ascii_letters
        for r in rows_to_update:
            token = "".join(random.choices(characters, k=10))
            cur.execute("UPDATE davetiyeler SET davetiye_kod = ? WHERE id = ?", (token, r[0]))
        con.commit()
        print(f"Guester: {len(rows_to_update)} adet eski davetiyeye rastgele kod atandı.")

    # Veritabanı göçü: Eğer olusturma_tarihi kolonu yoksa güvenli bir şekilde ekle
    cur.execute("PRAGMA table_info(davetiyeler)")
    columns = [c[1] for c in cur.fetchall()]
    if "olusturma_tarihi" not in columns:
        try:
            cur.execute("""
                CREATE TABLE davetiyeler_new(
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    gelin_adi TEXT,
                    damat_adi TEXT,
                    saat_bilgisi TEXT,
                    yer_bilgisi TEXT,
                    giris_sifresi TEXT,
                    davetiye_kod TEXT UNIQUE,
                    olusturma_tarihi TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            cur.execute("""
                INSERT INTO davetiyeler_new (id, gelin_adi, damat_adi, saat_bilgisi, yer_bilgisi, giris_sifresi, davetiye_kod, olusturma_tarihi)
                SELECT id, gelin_adi, damat_adi, saat_bilgisi, yer_bilgisi, giris_sifresi, davetiye_kod, CURRENT_TIMESTAMP 
                FROM davetiyeler
            """)
            cur.execute("DROP TABLE davetiyeler")
            cur.execute("ALTER TABLE davetiyeler_new RENAME TO davetiyeler")
            con.commit()
            print("Veritabanı başarıyla güncellendi (olusturma_tarihi sütunu eklendi).")
        except Exception as e:
            print(f"Veritabanı güncellenirken hata oluştu: {e}")
            con.rollback()
    # Veritabanı göçü: kullanici_id kolonu kontrolü ve eklenmesi
    cur.execute("PRAGMA table_info(davetiyeler)")
    columns = [c[1] for c in cur.fetchall()]
    if "kullanici_id" not in columns:
        try:
            print("Guester: kullanici_id kolonu ekleniyor...")
            cur.execute("ALTER TABLE davetiyeler ADD COLUMN kullanici_id INTEGER")
            
            # Mevcut tüm davetiyeleri ilk kayıtlı kullanıcıya (örneğin 'ceyda') atayalım
            cur.execute("SELECT id FROM kullanicilar ORDER BY id ASC LIMIT 1")
            first_user = cur.fetchone()
            if first_user:
                cur.execute("UPDATE davetiyeler SET kullanici_id = ?", (first_user[0],))
            con.commit()
            print("Veritabanı başarıyla güncellendi (kullanici_id sütunu eklendi).")
        except Exception as e:
            print(f"Hata (kullanici_id ekleme): {e}")
            con.rollback()
    else:
        con.commit()

    # Veritabanı göçü: kullanicilar tablosuna isletme_adi kolonu ekleme
    cur.execute("PRAGMA table_info(kullanicilar)")
    user_columns = [c[1] for c in cur.fetchall()]
    if "isletme_adi" not in user_columns:
        try:
            print("Guester: isletme_adi kolonu kullanicilar tablosuna ekleniyor...")
            cur.execute("ALTER TABLE kullanicilar ADD COLUMN isletme_adi TEXT")
            con.commit()
        except Exception as e:
            print(f"Hata (isletme_adi ekleme): {e}")
            con.rollback()

    con.close()

init_db()

# =========================
# SMTP AYARLARI
# =========================
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
SMTP_USER = "guester.platform@gmail.com"
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD") or os.environ.get("GMAIL_APP_PASSWORD") or "SIFREN"

# =========================
# GLOBAL DATA
# =========================
veri_listesi = []

# =========================
# HELPERS (YARDIMCI FONKSİYONLAR)
# =========================
def get_user(ad_soyad, sifre):
    con = db()
    cur = con.cursor()
    cur.execute("SELECT id, ad_soyad FROM kullanicilar WHERE ad_soyad=? AND sifre=?", (ad_soyad, sifre))
    user = cur.fetchone()
    con.close()
    return user

def user_exists(email):
    con = db()
    cur = con.cursor()
    cur.execute("SELECT id FROM kullanicilar WHERE email=?", (email,))
    data = cur.fetchone()
    con.close()
    return data is not None

def add_user(ad_soyad, sifre, email):
    con = db()
    cur = con.cursor()
    try:
        cur.execute(
            "INSERT INTO kullanicilar(ad_soyad, sifre, email) VALUES (?, ?, ?)",
            (ad_soyad, sifre, email)
        )
        con.commit()
        return True
    except:
        return False
    finally:
        con.close()

def find_user(email):
    con = db()
    cur = con.cursor()
    cur.execute("SELECT ad_soyad, sifre FROM kullanicilar WHERE email=?", (email,))
    user = cur.fetchone()
    con.close()
    return user

# Bilgisayarın yerel ağdaki IP adresini otomatik alan yardımcı fonksiyon
def get_local_ip():
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        # Herhangi bir adrese bağlanmaya çalışarak aktif yerel IP'yi alıyoruz
        s.connect(('10.255.255.255', 1))
        IP = s.getsockname()[0]
    except Exception:
        IP = '127.0.0.1'
    finally:
        s.close()
    return IP

# 1 Hafta Süresi Dolan Davetiyeleri ve Konukları Otomatik Silme
def temizle_eski_davetiyeler():
    con = db()
    cur = con.cursor()
    cur.execute("SELECT id FROM davetiyeler WHERE datetime(olusturma_tarihi) < datetime('now', '-7 days')")
    expired_ids = [r[0] for r in cur.fetchall()]
    if expired_ids:
        print(f"WEB APP: Süresi dolan (1 haftadan eski) {len(expired_ids)} etkinlik ve davetli listesi siliniyor...")
        for e_id in expired_ids:
            cur.execute("DELETE FROM konuklar WHERE davetiye_id = ?", (e_id,))
            cur.execute("DELETE FROM davetiyeler WHERE id = ?", (e_id,))
        con.commit()
    con.close()

# Word (.docx) Dosyasını Okuma Fonksiyonu
def oku_word(path):
    doc = docx.Document(path)
    yeni_veriler = []
    full_text = ""
    
    # 1. Tablolardaki metinleri oku
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join([cell.text for cell in row if cell.text]).strip()
            if row_text:
                full_text += row_text + "\n"
                
    # 2. Normal paragrafları oku
    for para in doc.paragraphs:
        if para.text.strip():
            full_text += para.text + "\n"
            
    # 3. Ayıklama işlemi
    for s in full_text.split("\n"):
        if gecerli_mi(s):
            i, m = ayir(s)
            if i and m:
                yeni_veriler.append({"isim": i, "masa": m})
                
    return yeni_veriler, full_text

# =========================
# OCR FİLTRELERİ (Sadece PDF/Word için)
# =========================
def gecerli_mi(s):
    s = str(s).lower().strip()
    return ("masa" in s or re.search(r'\d+', s))

def ayir(s):
    m = re.search(r'^(.*?)\s*[Mm]asa\s*(\d+)$', s)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return "", ""

# =========================
# ROUTES (ROTALAR)
# =========================
@app.route("/")
def index():
    return render_template("index.html", mesaj="", mesaj_tipi="")

@app.route("/yonetici")
def yonetici():
    if "admin_adi" not in session:
        return redirect("/")
    # Eski davetiyeleri temizle
    temizle_eski_davetiyeler()

    con = db()
    cur = con.cursor()
    
    # Kullanıcı detaylarını çek
    user_id = session.get("user_id")
    cur.execute("SELECT email, isletme_adi, ad_soyad FROM kullanicilar WHERE id = ?", (user_id,))
    u_row = cur.fetchone()
    user_email = u_row[0] if (u_row and u_row[0]) else ""
    user_isletme = u_row[1] if (u_row and u_row[1]) else ""
    if u_row:
        session["admin_adi"] = u_row[2]

    cur.execute("SELECT * FROM davetiyeler WHERE kullanici_id = ? ORDER BY id DESC", (user_id,))
    tum_davetiyeler = cur.fetchall()
    
    secili_id_param = request.args.get("davetiye_id")
    secili_davetiye_id = None
    
    if secili_id_param:
        try:
            secili_davetiye_id = int(secili_id_param)
        except ValueError:
            secili_davetiye_id = None
            
    # Güvenlik kontrolü: Seçili davetiye bu kullanıcıya mı ait?
    if secili_davetiye_id:
        cur.execute("SELECT id FROM davetiyeler WHERE id = ? AND kullanici_id = ?", (secili_davetiye_id, session.get("user_id")))
        if not cur.fetchone():
            secili_davetiye_id = None

    if not secili_davetiye_id and tum_davetiyeler:
        secili_davetiye_id = tum_davetiyeler[0][0] # En son oluşturulan
        
    son_konuklar = []
    etkinlik_adi = "Henüz bir etkinlik yok"
    
    davetiye_kod = ""
    if secili_davetiye_id:
        cur.execute("SELECT gelin_adi, damat_adi, davetiye_kod FROM davetiyeler WHERE id = ?", (secili_davetiye_id,))
        dav = cur.fetchone()
        if dav:
            etkinlik_adi = f"{dav[0]} & {dav[1]}"
            davetiye_kod = dav[2] or ""
        cur.execute("""
            SELECT id, isim, masa, durum 
            FROM konuklar 
            WHERE davetiye_id = ? 
            ORDER BY isim COLLATE NOCASE ASC
        """, (secili_davetiye_id,))
        son_konuklar = [{"id": r[0], "isim": r[1], "masa": r[2], "durum": r[3]} for r in cur.fetchall()]
        
    con.close()

    # Bilgisayarın yerel ağ IP'sini al
    yerel_ip = get_local_ip()
    
    # Bulut sunucuyu otomatik algıla (örneğin pythonanywhere'de tünele gerek yoktur)
    hostname = request.host.split(':')[0]
    is_local = hostname in ['localhost', '127.0.0.1'] or hostname.startswith('192.168.') or hostname.startswith('172.') or hostname.startswith('10.')
    
    if not is_local:
        h_url = request.host_url
        tunel_aktif = True
    else:
        h_url = g_public_url + "/" if g_public_url else ""
        tunel_aktif = (g_public_url is not None)

    return render_template(
        "yonetici.html",
        veri=son_konuklar,
        etkinlik_adi=etkinlik_adi,
        secili_davetiye_id=secili_davetiye_id,
        davetiye_kod=davetiye_kod,
        davetiyeler=tum_davetiyeler,
        admin_adi=session.get("admin_adi", "Yönetici"),
        toplam_konuk=len(son_konuklar),
        gelen_konuk=sum(1 for k in son_konuklar if k["durum"] == "Geldi"),
        yerel_ip=yerel_ip,
        host_url=h_url,
        tunel_aktif=tunel_aktif,
        user_email=user_email,
        user_isletme=user_isletme
    )

# =========================
# LOGIN / REGISTER / LOGOUT
# =========================
@app.route("/login", methods=["POST"])
def login():
    ad_soyad = request.form["ad_soyad"]
    sifre = request.form["sifre"]
    user = get_user(ad_soyad, sifre)
    if user:
        session["user_id"] = user[0]
        session["admin_adi"] = user[1]
        return redirect("/yonetici?login=true")
    return render_template("index.html", mesaj="Hatalı giriş", mesaj_tipi="error")

@app.route("/register", methods=["POST"])
def register():
    ad_soyad = request.form["ad_soyad"]
    sifre = request.form["sifre"]
    email = request.form["email"]

    # Şifre kuralları kontrolü (en az 8 karakter, en az 1 rakam)
    if len(sifre) < 8 or not any(c.isdigit() for c in sifre):
        return render_template("index.html", mesaj="Şifre en az 8 karakter olmalı ve en az 1 rakam içermelidir.", mesaj_tipi="error")

    if user_exists(email):
        return render_template("index.html", mesaj="Bu email zaten kayıtlı", mesaj_tipi="error")

    ok = add_user(ad_soyad, sifre, email)
    if ok:
        return render_template("index.html", mesaj="Kayıt başarılı", mesaj_tipi="success")
    return render_template("index.html", mesaj="Kayıt başarısız", mesaj_tipi="error")

@app.route("/logout")
def logout():
    session.clear()
    return redirect("/")

@app.route("/yonetici_guncelle", methods=["POST"])
def yonetici_guncelle():
    if "user_id" not in session:
        return redirect("/")
    
    ad_soyad = request.form.get("ad_soyad", "").strip()
    email = request.form.get("email", "").strip()
    isletme_adi = request.form.get("isletme_adi", "").strip()
    
    if not ad_soyad or not email:
        return redirect("/yonetici?error=missing_fields")
        
    con = db()
    cur = con.cursor()
    
    # Email benzersizlik kontrolü (başka bir kullanıcıya ait mi?)
    cur.execute("SELECT id FROM kullanicilar WHERE email = ? AND id != ?", (email, session.get("user_id")))
    existing = cur.fetchone()
    if existing:
        con.close()
        return redirect("/yonetici?error=email_exists")
        
    cur.execute("""
        UPDATE kullanicilar 
        SET ad_soyad = ?, email = ?, isletme_adi = ? 
        WHERE id = ?
    """, (ad_soyad, email, isletme_adi, session.get("user_id")))
    con.commit()
    con.close()
    
    session["admin_adi"] = ad_soyad
    return redirect("/yonetici?success=profile_updated")

@app.route("/yonetici_sifre_degistir", methods=["POST"])
def yonetici_sifre_degistir():
    if "user_id" not in session:
        return redirect("/")
        
    eski_sifre = request.form.get("eski_sifre", "")
    yeni_sifre = request.form.get("yeni_sifre", "")
    
    con = db()
    cur = con.cursor()
    cur.execute("SELECT sifre FROM kullanicilar WHERE id = ?", (session.get("user_id"),))
    row = cur.fetchone()
    
    if not row or row[0] != eski_sifre:
        con.close()
        return redirect("/yonetici?error=wrong_old_password")
        
    # Yeni şifre kuralları kontrolü
    if len(yeni_sifre) < 8 or not any(c.isdigit() for c in yeni_sifre):
        con.close()
        return redirect("/yonetici?error=weak_password")
        
    cur.execute("UPDATE kullanicilar SET sifre = ? WHERE id = ?", (yeni_sifre, session.get("user_id")))
    con.commit()
    con.close()
    
    return redirect("/yonetici?success=password_changed")

@app.route("/yonetici_hesap_sil", methods=["POST"])
def yonetici_hesap_sil():
    if "user_id" not in session:
        return redirect("/")
        
    user_id = session.get("user_id")
    con = db()
    cur = con.cursor()
    
    # 1. Bu kullanıcının tüm davetiyelerinin ID'lerini al
    cur.execute("SELECT id FROM davetiyeler WHERE kullanici_id = ?", (user_id,))
    davetiye_ids = [r[0] for r in cur.fetchall()]
    
    # 2. Bu davetiyelere ait tüm konukları sil
    for d_id in davetiye_ids:
        cur.execute("DELETE FROM konuklar WHERE davetiye_id = ?", (d_id,))
        
    # 3. Davetiyeleri sil
    cur.execute("DELETE FROM davetiyeler WHERE kullanici_id = ?", (user_id,))
    
    # 4. Şifre sıfırlama kodlarını sil
    cur.execute("SELECT email FROM kullanicilar WHERE id = ?", (user_id,))
    u_row = cur.fetchone()
    if u_row:
        cur.execute("DELETE FROM sifre_kodlari WHERE email = ?", (u_row[0],))
        
    # 5. Kullanıcıyı sil
    cur.execute("DELETE FROM kullanicilar WHERE id = ?", (user_id,))
    
    con.commit()
    con.close()
    
    session.clear()
    return render_template("index.html", mesaj="Hesabınız başarıyla silindi.", mesaj_tipi="success")

def send_email_via_http(to_email, subject, body, html_body=None):
    import urllib.request
    import json
    
    url = os.environ.get("EMAIL_SERVICE_URL")
    if not url or not url.startswith("http"):
        return False
        
    data = {
        "to": to_email,
        "subject": subject,
        "body": body
    }
    if html_body:
        data["htmlBody"] = html_body
        
    try:
        req = urllib.request.Request(
            url,
            data=json.dumps(data).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST"
        )
        with urllib.request.urlopen(req, timeout=10) as response:
            res_data = json.loads(response.read().decode("utf-8"))
            return res_data.get("status") == "success"
    except Exception as e:
        print(f"HTTP Email Hatası: {e}")
        return False

@app.route("/forgot-password", methods=["POST"])
def forgot_password():
    # Destek: Hem JSON hem de form verisi
    if request.is_json:
        data = request.get_json()
        email = data.get("email")
    else:
        email = request.form.get("email")
        
    email = email.strip() if email else ""
    user = find_user(email)
    if not user:
        return jsonify({"status": "error", "message": "Bu e-posta adresiyle kayıtlı bir kullanıcı bulunamadı."})

    ad_soyad, _ = user
    
    # 6 Haneli rastgele kod üret
    import random
    code = str(random.randint(100000, 999999))
    
    # Veritabanına kaydet
    con = db()
    cur = con.cursor()
    cur.execute("INSERT OR REPLACE INTO sifre_kodlari (email, kod, olusturma_tarihi) VALUES (?, ?, CURRENT_TIMESTAMP)", (email, code))
    con.commit()
    con.close()
    
    subject = "Şifre Sıfırlama Kodu"
    body = (
        f"Merhaba {ad_soyad},\n\n"
        f"Şifre sıfırlama talebiniz için tek kullanımlık doğrulama kodunuz: {code}\n\n"
        f"Bu kod 10 dakika geçerlidir. Şifrenizi sıfırlama ekranına girerek yeni şifrenizi belirleyebilirsiniz."
    )
    
    html_body = f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <style>
    body {{
      font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
      background-color: #f8fafc;
      color: #334155;
      margin: 0;
      padding: 0;
    }}
    .email-container {{
      max-width: 550px;
      margin: 40px auto;
      background: #ffffff;
      border-radius: 20px;
      overflow: hidden;
      box-shadow: 0 10px 30px rgba(0,0,0,0.05);
      border: 1px solid #e2e8f0;
    }}
    .header {{
      background: linear-gradient(135deg, #7c3aed 0%, #a855f7 100%);
      padding: 30px;
      text-align: center;
      color: #ffffff;
    }}
    .logo-container {{
      display: inline-block;
      width: 50px;
      height: 50px;
      background: rgba(255, 255, 255, 0.2);
      border-radius: 50%;
      text-align: center;
      line-height: 48px;
      font-size: 24px;
      font-weight: bold;
      color: #ffffff;
      border: 2px solid #ffffff;
      box-shadow: 0 4px 10px rgba(0,0,0,0.1);
      margin-bottom: 10px;
    }}
    .header h1 {{
      margin: 0;
      font-size: 20px;
      font-weight: 700;
      letter-spacing: 0.5px;
      color: #ffffff;
    }}
    .content {{
      padding: 40px;
      text-align: center;
    }}
    .welcome-text {{
      font-size: 16px;
      color: #64748b;
      margin-bottom: 25px;
      text-align: left;
    }}
    .main-message {{
      font-size: 15px;
      line-height: 1.6;
      color: #334155;
      margin-bottom: 30px;
      text-align: left;
    }}
    .code-box {{
      display: inline-block;
      background: #f3e8ff;
      border: 2px dashed #a855f7;
      color: #7c3aed;
      font-size: 32px;
      font-weight: 800;
      padding: 15px 40px;
      border-radius: 12px;
      letter-spacing: 4px;
      margin: 20px 0;
      box-shadow: 0 4px 15px rgba(124, 58, 237, 0.1);
    }}
    .warning-text {{
      font-size: 13px;
      color: #94a3b8;
      margin-top: 25px;
      border-top: 1px solid #f1f5f9;
      padding-top: 20px;
      text-align: left;
    }}
    .footer {{
      background-color: #f8fafc;
      padding: 20px;
      text-align: center;
      font-size: 12px;
      color: #94a3b8;
      border-top: 1px solid #e2e8f0;
    }}
  </style>
</head>
<body>
  <div class="email-container">
    <div class="header">
      <div class="logo-container">G</div>
      <h1>Guester Platform</h1>
    </div>
    <div class="content">
      <div class="welcome-text">Merhaba <strong>{ad_soyad}</strong>,</div>
      <div class="main-message">
        Şifre sıfırlama talebiniz için tek kullanımlık güvenlik kodunuz oluşturuldu. Aşağıdaki kodu kullanarak şifrenizi güvenle yenileyebilirsiniz:
      </div>
      <div class="code-box">{code}</div>
      <div class="warning-text">
        Bu kod güvenlik önlemleri sebebiyle <strong>10 dakika</strong> geçerlidir.<br>
        Talebi siz gerçekleştirmediyseniz bu e-postayı dikkate almayınız.
      </div>
    </div>
    <div class="footer">
      © 2026 Guester. Tüm hakları saklıdır.
    </div>
  </div>
</body>
</html>
"""

    # 1. Önce HTTP (Google Apps Script Web App) üzerinden göndermeyi dene
    if send_email_via_http(email, subject, body, html_body):
        return jsonify({"status": "success", "message": "E-postanıza 6 haneli doğrulama kodu gönderildi."})
        
    # 2. Eğer başarısız olursa veya tanımlı değilse klasik SMTP'yi (port 587) dene
    try:
        if SMTP_PASSWORD == "SIFREN" or not SMTP_PASSWORD:
            raise Exception("SMTP şifresi varsayılan değerde ('SIFREN')")
            
        msg = MIMEMultipart("alternative")
        msg["Subject"] = Header(subject, "utf-8")
        msg["From"] = SMTP_USER
        msg["To"] = email
        
        part1 = MIMEText(body, "plain", "utf-8")
        part2 = MIMEText(html_body, "html", "utf-8")
        msg.attach(part1)
        msg.attach(part2)
        
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SMTP_USER, SMTP_PASSWORD)
        server.sendmail(SMTP_USER, [email], msg.as_string())
        server.quit()
        return jsonify({"status": "success", "message": "E-postanıza 6 haneli doğrulama kodu gönderildi."})
    except Exception as e:
        print(f"SMTP Hatası: {e}")
        # Güvenlik için şifre sıfırlama kodunu artık ekrana YAZDIRMIYORUZ
        return jsonify({
            "status": "error",
            "message": "E-posta gönderilemedi. Lütfen sistem yöneticinizle veya .env ayarlarınızla iletişime geçin."
        })

@app.route("/verify-reset-code", methods=["POST"])
def verify_reset_code():
    data = request.get_json() or {}
    email = (data.get("email") or "").strip()
    code = (data.get("code") or "").strip()
    
    if not email or not code:
        return jsonify({"status": "error", "message": "Geçersiz parametreler."})
        
    con = db()
    cur = con.cursor()
    # 10 dakika içindeki kodları kabul et
    cur.execute("""
        SELECT kod FROM sifre_kodlari 
        WHERE email = ? AND datetime(olusturma_tarihi) >= datetime('now', '-10 minutes')
    """, (email,))
    row = cur.fetchone()
    con.close()
    
    if row and row[0] == code:
        return jsonify({"status": "success", "message": "Kod doğrulandı."})
    else:
        return jsonify({"status": "error", "message": "Hatalı veya süresi dolmuş kod girdiniz."})

@app.route("/reset-password", methods=["POST"])
def reset_password():
    data = request.get_json() or {}
    email = (data.get("email") or "").strip()
    code = (data.get("code") or "").strip()
    yeni_sifre = data.get("yeni_sifre")
    
    if not email or not code or not yeni_sifre or len(yeni_sifre) < 8 or not any(c.isdigit() for c in yeni_sifre):
        return jsonify({"status": "error", "message": "Yeni şifre en az 8 karakter olmalı ve en az 1 rakam içermelidir."})
        
    con = db()
    cur = con.cursor()
    
    # Son bir doğrulama yap
    cur.execute("""
        SELECT kod FROM sifre_kodlari 
        WHERE email = ? AND datetime(olusturma_tarihi) >= datetime('now', '-10 minutes')
    """, (email,))
    row = cur.fetchone()
    
    if row and row[0] == code:
        # Şifreyi güncelle
        cur.execute("UPDATE kullanicilar SET sifre = ? WHERE email = ?", (yeni_sifre, email))
        # Kodu sil
        cur.execute("DELETE FROM sifre_kodlari WHERE email = ?", (email,))
        con.commit()
        con.close()
        return jsonify({"status": "success", "message": "Şifreniz başarıyla sıfırlandı."})
    else:
        con.close()
        return jsonify({"status": "error", "message": "Geçersiz veya süresi dolmuş işlem."})

# ========================================================
# TEK TUŞLA ETKİNLİK VE LİSTE OLUŞTURMA ROTASI
# ========================================================
@app.route("/etkinlik_olustur", methods=["POST"])
def etkinlik_olustur():
    davetiye_hedef = request.form.get("davetiye_hedef", "yeni")
    files = request.files.getlist("dosya")
    
    con = db()
    cur = con.cursor()
    
    if davetiye_hedef == "yeni":
        gelin_adi = request.form.get("gelin_adi", "Gelin")
        damat_adi = request.form.get("damat_adi", "Damat")
        saat_bilgisi = request.form.get("saat_bilgisi", "Saat Belirtilmemiş")
        yer_bilgisi = request.form.get("yer_bilgisi", "Konum Belirtilmemiş")
        giris_sifresi = request.form.get("giris_sifresi", "1234")
        davetiye_kod = generate_unique_token()
        
        cur.execute("""
            INSERT INTO davetiyeler (gelin_adi, damat_adi, saat_bilgisi, yer_bilgisi, giris_sifresi, davetiye_kod, kullanici_id)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (gelin_adi, damat_adi, saat_bilgisi, yer_bilgisi, giris_sifresi, davetiye_kod, session.get("user_id")))
        davetiye_id = cur.lastrowid
    else:
        davetiye_id = int(davetiye_hedef)
        cur.execute("SELECT id FROM davetiyeler WHERE id = ? AND kullanici_id = ?", (davetiye_id, session.get("user_id")))
        if not cur.fetchone():
            con.close()
            return redirect("/yonetici")
        
    extracted_saat = None
    extracted_yer = None
    
    if files and not all(f.filename == '' for f in files):
        for file in files:
            if file.filename == '':
                continue
                
            path = os.path.join(UPLOAD_FOLDER, file.filename)
            try:
                file.save(path)
                
                yeni_veriler = []
                full_text = ""
                filename_lower = file.filename.lower()
                
                # IMAGE OKUMA (guester_ocr.py)
                if filename_lower.endswith((".png", ".jpg", ".jpeg")):
                    print(f"WEB APP: resmi_analiz_et cagrilacak file: {path} (filename: {file.filename})")
                    yeni_veriler = guester_ocr.resmi_analiz_et(path)
                    print(f"WEB APP: yeni_veriler okundu. Toplam: {len(yeni_veriler)} satir.")
                    full_text = "\n".join([f"{v['isim']} {v['masa']}" for v in yeni_veriler])
                
                # PDF OKUMA
                elif filename_lower.endswith(".pdf"):
                    with pdfplumber.open(path) as pdf:
                        for page in pdf.pages:
                            t = page.extract_text()
                            if t:
                                full_text += t + "\n"
                                for s in t.split("\n"):
                                    if gecerli_mi(s):
                                        i, m = ayir(s)
                                        if i and m:
                                            yeni_veriler.append({"isim": i, "masa": m})
                                            
                # WORD OKUMA (.docx)
                elif filename_lower.endswith(".docx"):
                    try:
                        yeni_veriler, doc_text = oku_word(path)
                        print(f"WEB APP: Word dosyasından {len(yeni_veriler)} davetli okundu.")
                        full_text = doc_text
                    except Exception as e:
                        print(f"WEB APP: Word okuma hatası: {e}")
                        
                # Okunan verileri doğrudan konuklar tablosuna ekle
                for item in yeni_veriler:
                    cur.execute("""
                        INSERT INTO konuklar (davetiye_id, isim, masa, durum)
                        VALUES (?, ?, ?, 'Gelmedi')
                    """, (davetiye_id, item["isim"], item["masa"]))
                    
                # Saat ve konum bilgisi ayıklama
                saat_match = re.search(r'(saat\s*|:\s*)(\d{2}[:\.]\d{2})', full_text, re.IGNORECASE)
                if saat_match:
                    extracted_saat = saat_match.group(2)
                    
                for satir in full_text.split("\n"):
                    if any(kelime in satir.lower() for kelime in ["salon", "akkm", "otel", "restoran", "bahçesi", "merkezi"]):
                        extracted_yer = satir.strip()
                        break
            except Exception as file_err:
                print(f"WEB APP HATA: {file.filename} işlenirken hata oluştu: {file_err}")
                
        # Eğer yeni davetiyeyse ve OCR'den saat/yer bulunduysa güncelle
        if davetiye_hedef == "yeni" and (extracted_saat or extracted_yer):
            if extracted_saat:
                cur.execute("UPDATE davetiyeler SET saat_bilgisi = ? WHERE id = ?", (extracted_saat, davetiye_id))
            if extracted_yer:
                cur.execute("UPDATE davetiyeler SET yer_bilgisi = ? WHERE id = ?", (extracted_yer, davetiye_id))
                
    con.commit()
    con.close()
    
    # Yönlendirmeden önce hostes sayfasına doğrudan erişim yetkisi verelim (link açıldığında hostese geçsin)
    session["giris_izni"] = True
    session["davetiye_id"] = davetiye_id
    
    # Seçili davetiyeyle birlikte yönetici sayfasına geri yönlendiriyoruz
    return redirect(f"/yonetici?davetiye_id={davetiye_id}")

# ========================================================
# AKTİF DAVETLİ LİSTESİ YÖNETİM METOTLARI (EKLE, SİL, GÜNCELLE, TEMİZLE)
# ========================================================
@app.route("/konuk_ekle", methods=["POST"])
def konuk_ekle():
    if "admin_adi" not in session:
        return redirect("/")
    davetiye_id = request.form.get("davetiye_id")
    isim = request.form.get("isim", "").strip()
    masa = request.form.get("masa", "").strip()
    
    if not davetiye_id or not isim:
        return redirect("/yonetici")
        
    con = db()
    cur = con.cursor()
    # Güvenlik kontrolü: Davetiye bu kullanıcıya mı ait?
    cur.execute("SELECT id FROM davetiyeler WHERE id = ? AND kullanici_id = ?", (int(davetiye_id), session.get("user_id")))
    if cur.fetchone():
        cur.execute("""
            INSERT INTO konuklar (davetiye_id, isim, masa, durum)
            VALUES (?, ?, ?, 'Gelmedi')
        """, (int(davetiye_id), isim, masa))
        con.commit()
    con.close()
    
    return redirect(f"/yonetici?davetiye_id={davetiye_id}")

@app.route("/konuk_sil/<int:id>")
def konuk_sil(id):
    if "admin_adi" not in session:
        return redirect("/")
    davetiye_id = request.args.get("davetiye_id", "")
    con = db()
    cur = con.cursor()
    # Güvenlik kontrolü: Silinecek konuk bu kullanıcının davetiyesine mi bağlı?
    cur.execute("""
        SELECT k.id FROM konuklar k
        JOIN davetiyeler d ON k.davetiye_id = d.id
        WHERE k.id = ? AND d.kullanici_id = ?
    """, (id, session.get("user_id")))
    if cur.fetchone():
        cur.execute("DELETE FROM konuklar WHERE id = ?", (id,))
        con.commit()
    con.close()
    if davetiye_id:
        return redirect(f"/yonetici?davetiye_id={davetiye_id}")
    return redirect("/yonetici")

@app.route("/konuk_guncelle/<int:id>")
def konuk_guncelle(id):
    if "admin_adi" not in session:
        return redirect("/")
    davetiye_id = request.args.get("davetiye_id", "")
    isim = request.args.get("isim", "").strip()
    masa = request.args.get("masa", "").strip()
    
    con = db()
    cur = con.cursor()
    # Güvenlik kontrolü: Güncellenecek konuk bu kullanıcının davetiyesine mi bağlı?
    cur.execute("""
        SELECT k.id FROM konuklar k
        JOIN davetiyeler d ON k.davetiye_id = d.id
        WHERE k.id = ? AND d.kullanici_id = ?
    """, (id, session.get("user_id")))
    if cur.fetchone():
        cur.execute("UPDATE konuklar SET isim = ?, masa = ? WHERE id = ?", (isim, masa, id))
        con.commit()
    con.close()
    if davetiye_id:
        return redirect(f"/yonetici?davetiye_id={davetiye_id}")
    return redirect("/yonetici")

@app.route("/konuklar_temizle/<int:davetiye_id>")
def konuklar_temizle(davetiye_id):
    if "admin_adi" not in session:
        return redirect("/")
    con = db()
    cur = con.cursor()
    # Güvenlik kontrolü
    cur.execute("SELECT id FROM davetiyeler WHERE id = ? AND kullanici_id = ?", (davetiye_id, session.get("user_id")))
    if cur.fetchone():
        cur.execute("DELETE FROM konuklar WHERE davetiye_id = ?", (davetiye_id,))
        con.commit()
    con.close()
    return redirect(f"/yonetici?davetiye_id={davetiye_id}")

@app.route("/davetiye_sil/<int:id>")
def davetiye_sil(id):
    if "admin_adi" not in session:
        return redirect("/")
    con = db()
    cur = con.cursor()
    # Güvenlik kontrolü
    cur.execute("SELECT id FROM davetiyeler WHERE id = ? AND kullanici_id = ?", (id, session.get("user_id")))
    if cur.fetchone():
        # Önce bu davetiyeye ait konukları sil
        cur.execute("DELETE FROM konuklar WHERE davetiye_id = ?", (id,))
        # Sonra davetiyenin kendisini sil
        cur.execute("DELETE FROM davetiyeler WHERE id = ?", (id,))
        con.commit()
    con.close()
    return redirect("/yonetici")

# =========================
# DINAMIK GUESTER GİRİŞ ROTASI
# =========================
@app.route("/inv/<string:davetiye_kod>", methods=["GET", "POST"])
def guester_login(davetiye_kod):
    hata_mesaji = None

    con = db()
    cur = con.cursor()
    
    # Sayısal bir değer girildiyse (örn: /inv/18) eski ID tabanlı sistemden sorgula
    if davetiye_kod.isdigit():
        cur.execute("SELECT * FROM davetiyeler WHERE id = ?", (int(davetiye_kod),))
    else:
        cur.execute("SELECT * FROM davetiyeler WHERE davetiye_kod = ?", (davetiye_kod,))
        
    davetiye = cur.fetchone()
    con.close()

    if not davetiye:
        return render_template("hata.html", kod=404, baslik="Sayfa Bulunamadı", aciklama="Aradığınız davetiye veya etkinlik sistemde kayıtlı değil ya da silinmiş olabilir."), 404

    davetiye_id = davetiye[0]
    etkinlik_verisi = {
        "id": davetiye_id,
        "gelin_adi": davetiye[1],
        "damat_adi": davetiye[2],
        "tarih": davetiye[3],  
        "konum": davetiye[4],
        "tur": f"{davetiye[1]} & {davetiye[2]} Düğün Organizasyonu"
    }
    dogru_sifre = davetiye[5]

    if request.method == "POST":
        sifre = request.form.get("sifre")
        if sifre == dogru_sifre:
            session["giris_izni"] = True
            session["davetiye_id"] = davetiye_id
            return redirect("/guester-panel")
        else:
            hata_mesaji = "Hatalı şifre girdiniz. Lütfen tekrar deneyin."

    return render_template("guester_login.html", etkinlik=etkinlik_verisi, hata=hata_mesaji)

@app.route("/guester-panel", methods=["GET", "POST"])
def guester_panel():
    if not session.get("giris_izni"):
        return render_template("hata.html", kod=403, baslik="Erişim Engellendi", aciklama="Bu panele erişebilmek için öncelikle davetiyenin giriş şifresini doğru girmelisiniz."), 403
        
    davetiye_id = session.get("davetiye_id")
    arama_sorgusu = ""
    if request.method == "POST":
        arama_sorgusu = request.form.get("isim", "").strip()
    elif request.method == "GET":
        arama_sorgusu = request.args.get("isim", "").strip()
        
    con = db()
    cur = con.cursor()
    
    # Etkinlik bilgilerini çek
    cur.execute("SELECT gelin_adi, damat_adi FROM davetiyeler WHERE id = ?", (davetiye_id,))
    davetiye = cur.fetchone()
    etkinlik_adi = f"{davetiye[0]} & {davetiye[1]}" if davetiye else "Düğün Organizasyonu"
    
    if arama_sorgusu:
        cur.execute("""
            SELECT id, isim, masa, durum 
            FROM konuklar 
            WHERE davetiye_id = ? AND isim LIKE ?
            ORDER BY id ASC
        """, (davetiye_id, f"%{arama_sorgusu}%"))
    else:
        cur.execute("""
            SELECT id, isim, masa, durum 
            FROM konuklar 
            WHERE davetiye_id = ?
            ORDER BY id ASC
        """, (davetiye_id,))
        
    konuklar = [{"id": r[0], "isim": r[1], "masa": r[2], "durum": r[3]} for r in cur.fetchall()]
    
    # Toplam ve gelen konuk istatistikleri
    cur.execute("SELECT COUNT(*), SUM(CASE WHEN durum = 'Geldi' THEN 1 ELSE 0 END) FROM konuklar WHERE davetiye_id = ?", (davetiye_id,))
    stats = cur.fetchone()
    toplam = stats[0] or 0
    gelen = stats[1] or 0
    
    con.close()
    
    return render_template(
        "guester.html", 
        veri=konuklar, 
        etkinlik_adi=etkinlik_adi, 
        arama_sorgusu=arama_sorgusu,
        toplam=toplam,
        gelen=gelen
    )

@app.route("/guester_durum_degis/<int:konuk_id>", methods=["POST"])
def guester_durum_degis(konuk_id):
    if not session.get("giris_izni"):
        return "Yetkisiz işlem! 🛑", 403
        
    con = db()
    cur = con.cursor()
    
    cur.execute("SELECT durum FROM konuklar WHERE id = ?", (konuk_id,))
    row = cur.fetchone()
    if row:
        yeni_durum = "Gelmedi" if row[0] == "Geldi" else "Geldi"
        cur.execute("UPDATE konuklar SET durum = ? WHERE id = ?", (yeni_durum, konuk_id))
        con.commit()
        
    con.close()
    return redirect("/guester-panel")

@app.route("/guester_durum_degis_ajax/<int:konuk_id>", methods=["POST"])
def guester_durum_degis_ajax(konuk_id):
    if not session.get("giris_izni"):
        return jsonify({"error": "Yetkisiz işlem"}), 403
        
    davetiye_id = session.get("davetiye_id")
    con = db()
    cur = con.cursor()
    
    cur.execute("SELECT durum FROM konuklar WHERE id = ?", (konuk_id,))
    row = cur.fetchone()
    yeni_durum = "Gelmedi"
    if row:
        yeni_durum = "Gelmedi" if row[0] == "Geldi" else "Geldi"
        cur.execute("UPDATE konuklar SET durum = ? WHERE id = ?", (yeni_durum, konuk_id))
        con.commit()
        
    # Güncel istatistikleri al
    cur.execute("SELECT COUNT(*), SUM(CASE WHEN durum = 'Geldi' THEN 1 ELSE 0 END) FROM konuklar WHERE davetiye_id = ?", (davetiye_id,))
    stats = cur.fetchone()
    toplam = stats[0] or 0
    gelen = stats[1] or 0
    
    con.close()
    return jsonify({
        "status": yeni_durum,
        "toplam": toplam,
        "gelen": gelen
    })

# =========================
# UYGULAMAYI ÇALIŞTIR
# =========================
if __name__ == "__main__":
    import sys
    if "--no-reloader" in sys.argv:
        log_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "bg_output.log")
        try:
            sys.stdout = open(log_path, "w", encoding="utf-8")
            sys.stderr = sys.stdout
        except:
            pass

    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    init_db()
    
    use_reloader = "--no-reloader" not in sys.argv
    
    # Sadece ana Werkzeug reloader sürecinde veya debug kapalıyken veya reloader devre dışıyken tüneli başlat
    if os.environ.get("WERKZEUG_RUN_MAIN") == "true" or not app.debug or not use_reloader:
        t = threading.Thread(target=start_ssh_tunnel, daemon=True)
        t.start()
        
    app.run(host="0.0.0.0", port=5000, debug=True, use_reloader=use_reloader)